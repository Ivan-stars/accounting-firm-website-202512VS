import re
import json
from docx import Document
from bs4 import BeautifulSoup

def extract_en_block(js_text):
    start = js_text.find("'en': {")
    if start == -1:
        print("Could not find 'en': { in lang.js")
        return None
    brace_count = 0
    in_block = False
    block_lines = []
    for line in js_text[start:].splitlines():
        if '{' in line:
            brace_count += line.count('{')
            in_block = True
        if in_block:
            block_lines.append(line)
        if '}' in line:
            brace_count -= line.count('}')
            if in_block and brace_count == 0:
                break
    block = '\n'.join(block_lines)
    block = re.sub(r'//.*', '', block)
    block = re.sub(r',\s*}', '}', block)
    block = re.sub(r',\s*\n\s*}', '\n}', block)
    block = block.replace("'", '"')
    block = re.sub(r'^"en":\s*', '', block)
    return block

with open('assets/js/lang.js', encoding='utf-8') as f:
    lang_js = f.read()

translations = {}
en_block = extract_en_block(lang_js)
if en_block:
    try:
        translations = json.loads(en_block)
    except Exception as e:
        print('Error parsing translations:', e)
        print('en_block:', en_block)
else:
    print('Could not extract en block from lang.js')

def get_i18n_text(key):
    return translations.get(key, f'[{key}]')

def is_chinese(text):
    return any('\u4e00' <= ch <= '\u9fff' for ch in text)

def clean_text(text):
    return re.sub(r'<[^>]+>', '', text)

def extract_visible_text(html):
    soup = BeautifulSoup(html, 'html.parser')
    print('--- Replacing data-i18n keys with English translations ---')
    for tag in soup.find_all(attrs={'data-i18n': True}):
        key = tag['data-i18n']
        en_val = get_i18n_text(key)
        print(f'data-i18n="{key}": "{en_val}"')
        tag.clear()
        tag.append(en_val)
    for s in soup(['script', 'style']):
        s.decompose()
    text_blocks = []
    readable_tags = ['h1', 'h2', 'h3', 'h4', 'h5', 'h6', 'p', 'li', 'a', 'address', 'button', 'span']
    print('--- Extracting text blocks from readable tags ---')
    for el in soup.find_all(readable_tags):
        t = el.get_text(strip=True)
        if t:
            t = clean_text(t)
            if is_chinese(t):
                found = False
                for k, v in translations.items():
                    if v == t:
                        text_blocks.append(translations[k])
                        found = True
                        break
                if not found:
                    text_blocks.append(f'[NEEDS TRANSLATION] {t}')
            else:
                text_blocks.append(t)
    print('--- Text blocks before deduplication ---')
    for tb in text_blocks:
        print(tb)
    seen = set()
    unique_blocks = []
    for x in text_blocks:
        x_stripped = x.strip()
        if not x_stripped:
            continue
        if x_stripped not in seen and not any(x_stripped in y for y in unique_blocks):
            unique_blocks.append(x_stripped)
            seen.add(x_stripped)
    print('--- Text blocks after deduplication ---')
    for ub in unique_blocks:
        print(ub)
    return unique_blocks

# Required Contact texts
required_texts = [
    "Contact",
    "Contact Information",
    "Get in Touch",
    "Office Address",
    "Business Hours"
]

doc = Document()
doc.add_heading('Raymond C W Tam and Co. Contact Content (English)', 0)

with open('contact.html', encoding='utf-8') as f:
    html = f.read()
text_blocks = extract_visible_text(html)
print('--- Writing to document ---')
for t in text_blocks:
    print(t)
    doc.add_paragraph(t)

# Check for required texts
missing = [txt for txt in required_texts if not any(txt in block for block in text_blocks)]
if missing:
    print('WARNING: The following required texts were NOT found in the export:')
    for m in missing:
        print('  -', m)
    print('Please check your HTML and translation files. The document will NOT be saved.')
else:
    doc.save('Contact.docx')
    print('Exported Contact page content to Contact.docx (all required texts found)') 