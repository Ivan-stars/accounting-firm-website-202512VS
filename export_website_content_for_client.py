import re
from docx import Document
from bs4 import BeautifulSoup
import jieba

# Helper to extract text from HTML with data-i18n keys
with open('assets/js/lang.js', encoding='utf-8') as f:
    lang_js = f.read()

# Extract English translations from lang.js
match = re.search(r"'en': \{([\s\S]+?)\}\s*// 导出函数", lang_js)
translations = {}
if match:
    en_block = match.group(1)
    for line in en_block.splitlines():
        m = re.match(r"\s*'([^']+)':\s*'([^']+)',?", line)
        if m:
            translations[m.group(1)] = m.group(2)

def get_i18n_text(key):
    return translations.get(key, f'[{key}]')

def is_chinese(text):
    # Check if the text contains Chinese characters
    return any('\u4e00' <= ch <= '\u9fff' for ch in text)

def extract_visible_text(html, lang='en'):
    soup = BeautifulSoup(html, 'html.parser')
    # Replace data-i18n with English translation
    for tag in soup.find_all(attrs={'data-i18n': True}):
        key = tag['data-i18n']
        tag.string = get_i18n_text(key)
    # Remove script/style
    for s in soup(['script', 'style']):
        s.decompose()
    # Get visible text
    text_blocks = []
    for el in soup.find_all(['h1','h2','h3','h4','h5','h6','p','li','a','address']):
        t = el.get_text(strip=True)
        if t:
            # Try to find translation if it's Chinese and not already translated
            if is_chinese(t):
                # Try to find a translation by value
                found = False
                for k, v in translations.items():
                    if v == t:
                        text_blocks.append(translations[k])
                        found = True
                        break
                if not found:
                    text_blocks.append(f'[NEEDS TRANSLATION] {t}')
            else:
                text_blocks.append(t)
    return text_blocks

doc = Document()
doc.add_heading('Raymond C W Tam and Co. Website Content (English)', 0)

pages = [
    ('Home Page', 'index.html'),
    ('About Page', 'about.html'),
    ('Services Page', 'services.html'),
    ('Contact Page', 'contact.html'),
]

for page_title, filename in pages:
    with open(filename, encoding='utf-8') as f:
        html = f.read()
    doc.add_page_break()
    doc.add_heading(page_title, level=1)
    text_blocks = extract_visible_text(html)
    for t in text_blocks:
        doc.add_paragraph(t)

doc.save('export_website_content_for_client.docx')
print('Exported English website content to export_website_content_for_client.docx (with untranslated content flagged)') 